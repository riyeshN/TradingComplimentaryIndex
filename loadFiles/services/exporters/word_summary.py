"""
RCA_Cij_Summary.docx generation.

Produces a Word document with a method section (variables table and Steps 1-3)
followed by one RCA/Cij year table per (reporter, primary-tier code), with the
US and China partner columns merged side by side. The method section adapts to
the scope's aggregation tier (HS4 heading / HS2 chapter / SITC section). Each
table reports three Cij forms per partner — DG sum, DG weighted-average, and
unweighted RCA product. Layout mirrors readings/RCA_Cij_Summary.docx.
"""

from __future__ import annotations

import logging
from pathlib import Path

import pandas as pd

from ._filter import tier_suffix

# Tier labels are provided by the caller (typically Scope.primary_labels).
# Source of canonical heading descriptions for the supported scopes:
# loadFiles/services/scope.py.

# 'K' (the aggregation tier) is inserted per-tier by _tier_vocabulary so the
# method section matches the scope being exported (HS4 / HS2 / SITC).
VARIABLES_TABLE = [
    ('i',    "Reporter country (e.g. Korea, Vietnam)"),
    ('j',    "Partner country (US or China)"),
    ('k',    "An HS 6-digit product"),
    ('X_ik', "Country i's export of product k to the world"),
    ('X_i',  "Country i's total exports to the world"),
    ('M_jk', "Country j's import of product k from the world"),
    ('M_j',  "Country j's total imports from the world"),
    ('T_k',  "World total exports of product k"),
    ('T',    "World total exports across all products"),
]


def _tier_vocabulary(tier_column: str) -> tuple[str, str, str]:
    """(tier_label, tier_noun, K_description) for the aggregation tier, so the
    method section reads correctly for whichever scope is exported."""
    return {
        'HS4':  ('HS4',  'heading', "An HS 4-digit heading (a group of HS6 products)"),
        'HS2':  ('HS2',  'chapter', "An HS 2-digit chapter (a group of HS6 products)"),
        'SITC': ('SITC', 'section', "An SITC Rev.4 section (a group of HS6 products)"),
    }.get(tier_column, ('HS4', 'heading', "An HS 4-digit heading (a group of HS6 products)"))


def _apply_page_layout(document) -> None:
    """Landscape + narrow margins so the wide multi-Cij tables have room."""
    from docx.enum.section import WD_ORIENT
    from docx.shared import Inches
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Inches(0.4)
    section.top_margin = section.bottom_margin = Inches(0.5)


def _style_table(table, header_pt: float = 8.5, body_pt: float = 8.5) -> None:
    """Bold the header row and shrink the font so many columns fit on the page."""
    from docx.shared import Pt
    for row_index, row in enumerate(table.rows):
        is_header = row_index == 0
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(header_pt if is_header else body_pt)
                    if is_header:
                        run.font.bold = True


def _format_number(value, decimals: int) -> str:
    if pd.isna(value):
        return ''
    return f"{value:.{decimals}f}"


def _normalise_tier_frame(frame: pd.DataFrame, tier_column: str) -> pd.DataFrame:
    """Rename tier-specific columns to scope-agnostic names so the doc builder
    can be tier-blind. Handles HS4 (`_4digit`), HS2 (`_2digit`) and SITC
    (`_sitc`) suffixes."""
    suffix = tier_suffix(tier_column)
    return frame.rename(columns={
        f'TCI_DG_{suffix}':            'TCI_DG',
        f'TCI_DG_WeightedAvg_{suffix}':'TCI_DG_WeightedAvg',
        f'TCI_RCA_Product_{suffix}':   'TCI_RCA_Product',
        f'RCA_Export_{suffix}':        'RCA_Export',
        f'RCA_Import_{suffix}':        'RCA_Import',
    })[['Country', 'Year', tier_column,
        'TCI_DG', 'TCI_DG_WeightedAvg', 'TCI_RCA_Product', 'RCA_Export', 'RCA_Import']]


def _merge_partner_frames(
    china: pd.DataFrame,
    us: pd.DataFrame,
    tier_column: str,
    countries: list[str] | None,
    tier_codes: list[str] | None,
) -> pd.DataFrame:
    china = _normalise_tier_frame(china, tier_column).copy()
    us    = _normalise_tier_frame(us,    tier_column).copy()
    if countries is not None:
        china = china[china['Country'].isin(countries)]
        us    = us[us['Country'].isin(countries)]
    if tier_codes is not None:
        china = china[china[tier_column].isin(tier_codes)]
        us    = us[us[tier_column].isin(tier_codes)]

    merged = china.merge(
        us, on=['Country', 'Year', tier_column], how='outer',
        suffixes=('_CN', '_US'),
    )
    # RCA Reporter Export is partner-independent and verified at pipeline time
    # by TCICalculator._verify_partner_invariants — no soft check needed here.

    merged['Year'] = pd.to_numeric(merged['Year'], errors='coerce')
    return merged.dropna(subset=['Year']).sort_values(['Country', tier_column, 'Year'])


def _add_method_section(document, tier_column: str = 'HS4') -> None:
    tier_label, tier_noun, k_description = _tier_vocabulary(tier_column)

    document.add_heading('Method', level=1)

    document.add_heading('Variables', level=2)
    table = document.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = 'Symbol'
    table.rows[0].cells[1].text = 'Meaning'
    # Insert the aggregation-tier symbol K (tier-specific) right after k.
    variables = list(VARIABLES_TABLE)
    variables.insert(3, ('K', k_description))
    for symbol, meaning in variables:
        row = table.add_row().cells
        row[0].text = symbol
        row[1].text = meaning
    document.add_paragraph('All trade values in USD thousands. Source: ITC Trade Map.')

    document.add_heading('Step 1 — Balassa RCA at HS6', level=2)
    document.add_paragraph(
        "How specialised is country i in exporting product k, compared with the world average?")
    document.add_paragraph("RCA_export = (X_ik / X_i) ÷ (T_k / T)")
    document.add_paragraph("RCA_import = (M_jk / M_j) ÷ (T_k / T)")

    document.add_heading('Step 2 — Drysdale–Garnaut Cij at HS6', level=2)
    document.add_paragraph(
        "Do reporter i's exports of product k line up with partner j's imports of k?")
    document.add_paragraph("Cij = (X_ik / X_i) × (M_jk / M_j) × (T / T_k)")
    document.add_paragraph("Equivalently, using the two RCAs from Step 1:")
    document.add_paragraph("Cij = RCA_export × RCA_import × (T_k / T)")

    document.add_heading(f'Step 3 — Aggregate from HS6 up to {tier_label}', level=2)
    document.add_paragraph(
        f"RCA at {tier_label} is computed by adding up the trade values across all "
        f"HS6 codes inside the {tier_noun} first, then applying the Balassa formula "
        f"at the {tier_label} level. So the formula is the same as Step 1, just with "
        f"the {tier_noun}-level totals as inputs.")
    document.add_paragraph(
        f"Cij at {tier_label} is the sum of all the HS6 Cij values inside the "
        f"{tier_noun}. No second weighting:")
    document.add_paragraph(
        f"Cij at {tier_label} ({tier_noun} K) = sum of Cij(k) for every HS6 code k inside K")
    document.add_paragraph(
        "Column key — three tier-level Cij forms are reported side by side (each "
        "shown for the US and for China):")
    document.add_paragraph(
        "  • \"Cij sum\" — the sum of HS6 Drysdale–Garnaut values above, each "
        "carrying its own world-share weight (T_k / T). Additive: sums to the "
        "headline index, equals Yang's comprehensive aggregate.")
    document.add_paragraph(
        f"  • \"Cij wtd-avg\" — the world-trade-share weighted mean of the "
        f"per-product complementarities: Σ_k (T_k / T_K) × RCA_export_k × "
        f"RCA_import_k, where T_K is the {tier_noun}-level world total and the "
        f"weights T_k/T_K sum to 1. Equivalently the DG sum scaled by T/T_K. "
        f"Same scale as an HS6 Cij, so it is comparable across {tier_noun}s; "
        f"not additive to the headline.")
    document.add_paragraph(
        f"  • \"Cij RCA×RCA\" — the {tier_noun}-level RCA_export × RCA_import, "
        "with no world-share weight (an unweighted product of the two "
        "specialisation ratios). The per-product index published by Yang "
        "(2023, Table 3).")

    document.add_heading('Data notes', level=2)
    if tier_column == 'HS2':
        document.add_paragraph(
            "Strategic scope covers 10 of the 11 HS2 chapters in the Freund, "
            "Mattoo, Mulabdic & Ruta (JIE) strategic-industry list. Chapter 98 is "
            "omitted: HS chapters 98–99 are reserved for national use and are not "
            "part of the internationally harmonized HS nomenclature (chapters "
            "1–97), so they carry no comparable trade data in ITC Trade Map or "
            "UN Comtrade.")
    document.add_paragraph(
        "HS code 8524 is reported for 2022 onward only. The code was vacated in "
        "the HS 2007 revision (where it had meant recorded media) and "
        "reintroduced in HS 2022 with a new meaning (flat-panel display "
        "modules). During the vacated years (2007–2021) world trade in the code "
        "is near zero, so any residual reporting produces meaningless comparative-"
        "advantage ratios; those years are excluded so the 8524 series reflects a "
        "single, current product definition.")
    document.add_paragraph(
        "Each table reports three Cij columns per partner: the DG-sum (additive, "
        "headline-consistent; recommended primary), the DG weighted-average "
        "(same scale across headings; not additive), and the unweighted "
        "RCA-product (matching Yang 2023). The RCA-product can take large values "
        "when both the reporter and partner are highly specialised in a narrow "
        "high-value heading; such values are genuine, not errors.")


def _add_reporter_tier_table(
    document, reporter: str, tier_code: str, label: str, year_rows: pd.DataFrame,
    tier_prefix: str = 'HS',
) -> None:
    heading = f"{reporter} — {tier_prefix} {tier_code}"
    if label:
        heading += f" ({label})"
    document.add_heading(heading, level=2)

    table = document.add_table(rows=1, cols=10)
    table.style = 'Light Grid Accent 1'
    # Grouped, plain-language headers. Cij forms: "sum" = DG sum of HS6 (primary),
    # "wtd-avg" = world-share weighted average, "RCA×RCA" = unweighted RCA product.
    headers = ['Year',
               'RCA export', 'RCA import\n(US)', 'RCA import\n(China)',
               'Cij sum\n(US)', 'Cij sum\n(China)',
               'Cij wtd-avg\n(US)', 'Cij wtd-avg\n(China)',
               'Cij RCA×RCA\n(US)', 'Cij RCA×RCA\n(China)']
    for col, text in enumerate(headers):
        table.rows[0].cells[col].text = text

    for _, year_row in year_rows.iterrows():
        rca_export = year_row['RCA_Export_CN']
        if pd.isna(rca_export):
            rca_export = year_row['RCA_Export_US']
        cells = table.add_row().cells
        cells[0].text = str(int(year_row['Year']))
        cells[1].text = _format_number(rca_export, 3)
        cells[2].text = _format_number(year_row['RCA_Import_US'], 3)
        cells[3].text = _format_number(year_row['RCA_Import_CN'], 3)
        cells[4].text = _format_number(year_row['TCI_DG_US'], 4)
        cells[5].text = _format_number(year_row['TCI_DG_CN'], 4)
        cells[6].text = _format_number(year_row['TCI_DG_WeightedAvg_US'], 4)
        cells[7].text = _format_number(year_row['TCI_DG_WeightedAvg_CN'], 4)
        cells[8].text = _format_number(year_row['TCI_RCA_Product_US'], 4)
        cells[9].text = _format_number(year_row['TCI_RCA_Product_CN'], 4)
    _style_table(table)


def _add_single_partner_tier_table(
    document, reporter: str, partner: str, tier_code: str, label: str,
    year_rows: pd.DataFrame, tier_column: str, tier_prefix: str,
) -> None:
    """One table per (reporter, tier-code) for a single reporter→partner pair.
    Columns: Year | RCA export | RCA import | Cij sum | Cij wtd-avg | Cij RCA×RCA."""
    suffix = tier_suffix(tier_column)
    heading = f"{reporter} → {partner} — {tier_prefix} {tier_code}"
    if label:
        heading += f" ({label})"
    document.add_heading(heading, level=2)

    table = document.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'
    headers = ['Year', 'RCA export', 'RCA import',
               'Cij sum', 'Cij wtd-avg', 'Cij RCA×RCA']
    for col, text in enumerate(headers):
        table.rows[0].cells[col].text = text

    for _, row in year_rows.sort_values('Year').iterrows():
        cells = table.add_row().cells
        cells[0].text = str(int(row['Year']))
        cells[1].text = _format_number(row[f'RCA_Export_{suffix}'], 3)
        cells[2].text = _format_number(row[f'RCA_Import_{suffix}'], 3)
        cells[3].text = _format_number(row[f'TCI_DG_{suffix}'], 4)
        cells[4].text = _format_number(row[f'TCI_DG_WeightedAvg_{suffix}'], 4)
        cells[5].text = _format_number(row[f'TCI_RCA_Product_{suffix}'], 4)
    _style_table(table)


def export_word_summary(
    tier_index_by_partner: dict[str, pd.DataFrame],
    export_dir: Path,
    tier_column: str = 'HS4',
    tier_labels: dict[str, str] | None = None,
    countries: list[str] | None = None,
    tier_codes: list[str] | None = None,
    logger: logging.Logger | None = None,
) -> Path | None:
    """
    Write RCA_Cij_Summary.docx to export_dir.

    Two layouts:
      - Two partners present (China + US): one table per (reporter, tier-code)
        with merged US/CN columns (ICT, strategic scopes).
      - A single partner (e.g. CCE for the Yang scope): one table per
        (reporter, tier-code) with RCA export, RCA import, and both Cij forms.

    Returns the output path, or None if no usable partner frames are present.
    """
    log = logger or logging.getLogger(__name__)
    tier_labels = tier_labels or {}
    tier_prefix = 'SITC' if tier_column == 'SITC' else 'HS'

    from docx import Document

    china = tier_index_by_partner.get('China')
    us = tier_index_by_partner.get('US')

    # ── Two-partner layout (ICT / strategic) ─────────────────────────────────
    if china is not None and us is not None:
        merged = _merge_partner_frames(china, us, tier_column, countries, tier_codes)
        document = Document()
        _apply_page_layout(document)
        document.add_heading(
            'RCA and Cij — Indo-Pacific reporters, US and China as partners, 2001–2024',
            level=0,
        )
        _add_method_section(document, tier_column)
        table_count = 0
        for reporter in sorted(merged['Country'].unique()):
            document.add_heading(str(reporter), level=1)
            reporter_rows = merged[merged['Country'] == reporter]
            for tier_code in sorted(reporter_rows[tier_column].unique()):
                year_rows = reporter_rows[reporter_rows[tier_column] == tier_code]
                _add_reporter_tier_table(
                    document, reporter, tier_code, tier_labels.get(tier_code, ''),
                    year_rows.sort_values('Year'), tier_prefix,
                )
                table_count += 1
        export_dir.mkdir(parents=True, exist_ok=True)
        output_path = export_dir / 'RCA_Cij_Summary.docx'
        document.save(output_path)
        log.info("Exported %s (%d reporter-%s tables).", output_path, table_count, tier_column)
        return output_path

    # ── Single-partner layout (e.g. Yang china→CCE) ──────────────────────────
    if len(tier_index_by_partner) != 1:
        log.warning("Word summary skipped: expected China+US or a single partner.")
        return None
    partner_name, frame = next(iter(tier_index_by_partner.items()))
    frame = frame.copy()
    if countries is not None:
        frame = frame[frame['Country'].isin(countries)]
    if tier_codes is not None:
        frame = frame[frame[tier_column].isin(tier_codes)]
    frame['Year'] = pd.to_numeric(frame['Year'], errors='coerce')
    frame = frame.dropna(subset=['Year'])

    document = Document()
    _apply_page_layout(document)
    document.add_heading(
        f'RCA and Cij — exporter vs {partner_name} (importer), by {tier_prefix} section',
        level=0,
    )
    _add_method_section(document, tier_column)
    table_count = 0
    for reporter in sorted(frame['Country'].unique()):
        document.add_heading(str(reporter), level=1)
        reporter_rows = frame[frame['Country'] == reporter]
        for tier_code in sorted(reporter_rows[tier_column].unique()):
            year_rows = reporter_rows[reporter_rows[tier_column] == tier_code]
            _add_single_partner_tier_table(
                document, reporter, partner_name, tier_code,
                tier_labels.get(tier_code, ''), year_rows, tier_column, tier_prefix,
            )
            table_count += 1

    export_dir.mkdir(parents=True, exist_ok=True)
    output_path = export_dir / 'RCA_Cij_Summary.docx'
    document.save(output_path)
    log.info("Exported %s (%d %s tables, single partner %s).",
             output_path, table_count, tier_column, partner_name)
    return output_path
