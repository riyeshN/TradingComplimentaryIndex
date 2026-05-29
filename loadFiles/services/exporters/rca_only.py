"""
RCA-only deliverable (strategic / HS2 scope).

Deliverables (all inside one folder, default `export/strategic/rca_tables/`):
  - `RCA_only.docx` — two tables:
      1) RCA import index for the US and China across the 10 strategic HS2
         chapters, one sub-table per country (rows=year, cols=sector with
         `n=<HS6 count>` headers). RCA import is reporter-independent, so
         reporter identity is not a dimension of the table.
      2) RCA import index for the US and China for HS 8541 and 8542
         (semiconductors and integrated circuits), rows=year, cols=country×HS4.
  - `RCA_import_evolution_US.png` and `RCA_import_evolution_China.png` —
     one line per strategic HS2 chapter, x-axis = year (2001–2024).

Inputs: `hs2_index_by_partner` and `hs4_index_by_partner` from
`TCICalculator(scope=SCOPE_STRATEGIC)`.
"""

from __future__ import annotations

import logging
from pathlib import Path

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns

from .word_summary import _apply_page_layout, _format_number, _style_table

STRATEGIC_CHAPTERS = ('28', '29', '30', '38', '84', '85', '87', '88', '90', '93')
CHAPTER_LABELS = {
    '28': 'Inorganic chem.',
    '29': 'Organic chem.',
    '30': 'Pharmaceuticals',
    '38': 'Misc. chemicals',
    '84': 'Machinery',
    '85': 'Electrical mach.',
    '87': 'Vehicles',
    '88': 'Aircraft',
    '90': 'Instruments',
    '93': 'Arms / ammunition',
}
PARTNERS = ('US', 'China')
SEMI_HS4 = ('8541', '8542')
YEAR_RANGE = (2001, 2024)
DATA_SOURCE_TEXT = 'Data source: Calculated from the World Trade Map.'


def _import_rca_by_year_sector(
    hs2_index_by_partner: dict[str, pd.DataFrame],
    partner: str,
) -> pd.DataFrame:
    """Pivot to rows=Year, cols=HS2 chapter, values=RCA_Import_2digit. RCA
    import is reporter-independent — collapse via mean (all reporters carry
    the same value)."""
    frame = hs2_index_by_partner[partner]
    frame = frame[frame['HS2'].isin(STRATEGIC_CHAPTERS)].copy()
    frame['Year'] = pd.to_numeric(frame['Year'], errors='coerce')
    frame = frame.dropna(subset=['Year'])
    frame = frame[(frame['Year'] >= YEAR_RANGE[0]) & (frame['Year'] <= YEAR_RANGE[1])]
    pivoted = (
        frame.groupby(['Year', 'HS2'])['RCA_Import_2digit']
        .mean()
        .unstack('HS2')
        .reindex(columns=STRATEGIC_CHAPTERS)
        .sort_index()
    )
    pivoted.index = pivoted.index.astype(int)
    return pivoted


def _hs6_count_per_chapter(
    hs6_with_indicators_by_partner: dict[str, pd.DataFrame],
    partner: str,
) -> dict[str, int]:
    """Distinct HS6 product codes per HS2 chapter for the given partner — the
    number of HS6 lines that aggregate up into each chapter."""
    frame = hs6_with_indicators_by_partner[partner].copy()
    frame['HS6'] = frame['Product code'].astype(str)
    frame['HS2'] = frame['HS6'].str[:2]
    frame = frame[frame['HS2'].isin(STRATEGIC_CHAPTERS)]
    return {
        chapter: int(frame.loc[frame['HS2'] == chapter, 'HS6'].nunique())
        for chapter in STRATEGIC_CHAPTERS
    }


def _semi_rca_table(hs4_index_by_partner: dict[str, pd.DataFrame]) -> pd.DataFrame:
    """Rows=Year, cols=MultiIndex(partner, HS4) for HS 8541 and 8542 import
    RCA. Reporter-independent — collapse via mean across reporters."""
    parts = []
    for partner in PARTNERS:
        frame = hs4_index_by_partner[partner]
        sub = frame[frame['HS4'].isin(SEMI_HS4)].copy()
        sub['Year'] = pd.to_numeric(sub['Year'], errors='coerce')
        sub = sub.dropna(subset=['Year'])
        sub = sub[(sub['Year'] >= YEAR_RANGE[0]) & (sub['Year'] <= YEAR_RANGE[1])]
        pivoted = (
            sub.groupby(['Year', 'HS4'])['RCA_Import_4digit']
            .mean()
            .unstack('HS4')
            .reindex(columns=SEMI_HS4)
        )
        pivoted.columns = pd.MultiIndex.from_product([[partner], SEMI_HS4])
        parts.append(pivoted)
    combined = pd.concat(parts, axis=1).sort_index()
    combined.index = combined.index.astype(int)
    return combined


# ── Word doc ────────────────────────────────────────────────────────────────

def _add_import_rca_table(document, partner: str,
                          pivoted: pd.DataFrame,
                          hs6_counts: dict[str, int]) -> None:
    document.add_heading(partner, level=2)
    n_cols = 1 + len(STRATEGIC_CHAPTERS)
    table = document.add_table(rows=1, cols=n_cols)
    table.style = 'Light Grid Accent 1'
    header = table.rows[0].cells
    header[0].text = 'Year'
    for col_index, chapter in enumerate(STRATEGIC_CHAPTERS, start=1):
        label = CHAPTER_LABELS.get(chapter, '')
        header[col_index].text = f"HS {chapter}\n{label}\n(n={hs6_counts.get(chapter, 0)})"
    for year, row in pivoted.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(int(year))
        for col_index, chapter in enumerate(STRATEGIC_CHAPTERS, start=1):
            cells[col_index].text = _format_number(row.get(chapter), 3)
    _style_table(table)


def _add_semi_table(document, semi_rca: pd.DataFrame) -> None:
    columns = list(semi_rca.columns)  # [(US,8541),(US,8542),(China,8541),(China,8542)]
    table = document.add_table(rows=1, cols=1 + len(columns))
    table.style = 'Light Grid Accent 1'
    header = table.rows[0].cells
    header[0].text = 'Year'
    for col_index, (partner, hs4) in enumerate(columns, start=1):
        header[col_index].text = f"{partner}\nHS {hs4}"
    for year, row in semi_rca.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(int(year))
        for col_index, key in enumerate(columns, start=1):
            cells[col_index].text = _format_number(row.get(key), 3)
    _style_table(table)


def _write_word(
    hs2_index_by_partner: dict[str, pd.DataFrame],
    hs4_index_by_partner: dict[str, pd.DataFrame],
    hs6_with_indicators_by_partner: dict[str, pd.DataFrame],
    export_dir: Path,
    log: logging.Logger,
) -> Path:
    from docx import Document

    document = Document()
    _apply_page_layout(document)

    document.add_heading(
        'RCA index for import (Comparative disadvantage) for the US and China '
        'in 10 Strategic Sectors, 2001-2024',
        level=1,
    )
    document.add_paragraph(
        "RCA values for imports of the 10 strategic HS2 chapters by the US and "
        "China. The RCA import index is reporter-independent (calculated from "
        "the importer's own world-trade shares), so the same value applies to "
        "every Indo-Pacific reporter and is reported once per (year, sector, "
        "importer). The `n` figure in each column header is the number of HS6 "
        "product lines that aggregate up into that HS2 chapter."
    )
    for partner in PARTNERS:
        pivoted = _import_rca_by_year_sector(hs2_index_by_partner, partner)
        counts = _hs6_count_per_chapter(hs6_with_indicators_by_partner, partner)
        _add_import_rca_table(document, partner, pivoted, counts)
    document.add_paragraph(DATA_SOURCE_TEXT)

    document.add_heading(
        'RCA index for import for the US and China in HS 8541 and 8542, 2001-2024',
        level=1,
    )
    document.add_paragraph(
        "Bilateral semiconductor headings: HS 8541 (diodes, transistors, "
        "photovoltaic cells) and HS 8542 (electronic integrated circuits). "
        "RCA import values are reporter-independent."
    )
    semi_rca = _semi_rca_table(hs4_index_by_partner)
    _add_semi_table(document, semi_rca)
    document.add_paragraph(DATA_SOURCE_TEXT)

    output_path = export_dir / 'RCA_only.docx'
    document.save(output_path)
    log.info("Wrote %s (import RCA tables + HS 8541/8542).", output_path)
    return output_path


# ── Charts ──────────────────────────────────────────────────────────────────

def _chapter_legend_label(code: str) -> str:
    return f"HS {code} — {CHAPTER_LABELS.get(code, '')}"


def _plot_partner_import_evolution(
    pivoted: pd.DataFrame,
    partner: str,
    output_path: Path,
) -> None:
    long = (
        pivoted.reset_index()
        .melt(id_vars='Year', var_name='HS2', value_name='RCA')
        .dropna(subset=['RCA'])
    )
    long['Chapter'] = long['HS2'].map(_chapter_legend_label)
    long = long.sort_values(['HS2', 'Year'])

    fig, ax = plt.subplots(figsize=(12, 6))
    sns.lineplot(
        data=long, x='Year', y='RCA', hue='Chapter',
        palette='tab10', marker='o', ax=ax,
    )
    ax.axhline(1, linestyle='--', color='grey', alpha=0.6, linewidth=1)
    ax.set_title(f"RCA index for {partner} import in 10 strategic sectors, 2001-2024")
    ax.set_xlabel('Year')
    ax.set_ylabel('RCA (import)')
    ax.set_xticks(range(YEAR_RANGE[0], YEAR_RANGE[1] + 1, 2))
    ax.set_xlim(YEAR_RANGE[0], YEAR_RANGE[1])
    ax.tick_params(axis='x', rotation=45)
    ax.legend(title='HS2 sector', loc='center left',
              bbox_to_anchor=(1, 0.5), fontsize=8)
    fig.text(0.01, 0.01, DATA_SOURCE_TEXT, fontsize=8, style='italic')
    fig.tight_layout(rect=[0, 0.03, 1, 1])
    fig.savefig(output_path, dpi=130, bbox_inches='tight')
    plt.close(fig)


def _plot_evolutions(
    hs2_index_by_partner: dict[str, pd.DataFrame],
    export_dir: Path,
    log: logging.Logger,
) -> int:
    count = 0
    for partner in PARTNERS:
        pivoted = _import_rca_by_year_sector(hs2_index_by_partner, partner)
        _plot_partner_import_evolution(
            pivoted, partner,
            export_dir / f"RCA_import_evolution_{partner}.png",
        )
        count += 1
    log.info("Wrote %d partner import-RCA evolution charts.", count)
    return count


# ── Public entry point ──────────────────────────────────────────────────────

def export_rca_only(
    hs2_index_by_partner: dict[str, pd.DataFrame],
    hs4_index_by_partner: dict[str, pd.DataFrame],
    hs6_with_indicators_by_partner: dict[str, pd.DataFrame],
    export_dir: Path,
    logger: logging.Logger | None = None,
) -> Path:
    """Write the RCA-only deliverable into `export_dir`. Returns the Word doc path."""
    log = logger or logging.getLogger(__name__)
    for name, mapping in (('hs2_index_by_partner', hs2_index_by_partner),
                          ('hs4_index_by_partner', hs4_index_by_partner),
                          ('hs6_with_indicators_by_partner', hs6_with_indicators_by_partner)):
        missing = [p for p in PARTNERS if p not in mapping]
        if missing:
            raise ValueError(f"{name} missing partner(s): {missing}")
    export_dir.mkdir(parents=True, exist_ok=True)

    sns.set_theme(style='whitegrid', context='notebook')
    word_path = _write_word(
        hs2_index_by_partner, hs4_index_by_partner,
        hs6_with_indicators_by_partner, export_dir, log,
    )
    _plot_evolutions(hs2_index_by_partner, export_dir, log)
    return word_path
