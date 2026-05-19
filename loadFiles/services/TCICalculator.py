import numpy as np
import pandas as pd
from pathlib import Path
import logging
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

EXPORT_DIR = "./data/TradeMapData/export/"
# UNCTAD ICT goods classification (HS 2022) — 23 HS4 headings, 107 HS6 codes.
# Source: readings/ICTList.docx
HS4_ICT_HEADINGS = [
    '8443', '8470', '8471', '8472', '8473',
    '8517', '8518', '8519', '8521', '8522', '8523', '8524', '8525',
    '8527', '8528', '8529', '8531', '8534', '8540', '8541', '8542',
    '9013', '9504',
]
# HS4 heading names are classification metadata, not trade data — the DB stores
# only HS6 product labels (HSProduct), so heading-level names are not available
# there. These come from the same source as HS4_ICT_HEADINGS: the UNCTAD ICT goods
# classification / ITC Trade Map (readings/ICTList.docx).
HS4_ICT_LABELS = {
    '8443': 'Printing machinery',
    '8470': 'Calculating machines, data recording',
    '8471': 'Automatic data-processing machines',
    '8472': 'Office machines',
    '8473': 'Parts and accessories',
    '8517': 'Telephone sets, smartphones, network apparatus',
    '8518': 'Microphones, loudspeakers',
    '8519': 'Sound recording / reproducing apparatus',
    '8521': 'Video recording / reproducing apparatus',
    '8522': 'Parts for sound and video apparatus',
    '8523': 'Discs, tapes, solid-state storage, smart cards',
    '8524': 'Flat panel display modules',
    '8525': 'Transmission apparatus for radio / TV',
    '8527': 'Reception apparatus for radio-broadcasting',
    '8528': 'Monitors, projectors',
    '8529': 'Parts for displays and transmission apparatus',
    '8531': 'Electric signalling apparatus',
    '8534': 'Printed circuits',
    '8540': 'Thermionic, cathode and photo-cathode tubes',
    '8541': 'Semiconductor devices, photosensitive elements',
    '8542': 'Electronic integrated circuits',
    '9013': 'Lasers, optical appliances',
    '9504': 'Video game consoles',
}
YEARS_TO_INCLUDE = [str(y) for y in range(2001, 2025)]


class TCICalculator:
    """
    Calculates Trade Complementarity Index (TCI) over the UNCTAD ICT goods scope
    (HS 2022, 23 HS4 headings, 107 HS6 codes) for Indo-Pacific reporters versus
    China and the United States.

    Three-tier output, all derived from a single HS6 source:
      - HS6 Cij (Drysdale-Garnaut)         — single-product values, audit trail.
      - HS4 Cij  = sum of HS6 Cij in heading — preserves bilateral product matching.
      - Headline = sum of HS6 Cij over scope — single Cij per (reporter, partner, year).

    Two TCI formulations are reported at every tier:
      - Drysdale-Garnaut: (Xi_k/Xi) × (Mj_k/Mj) × (WX/WX_k)              — primary.
      - Tan Fen (2024):   RCA_export × RCA_import (no world-share factor) — secondary.

    HS4 RCA is reported as an auxiliary heading-level specialisation metric. It is
    derived as a weighted average of HS6 RCA values with weight = (W_k / W_HS4),
    and is not used in the HS4 Cij calculation.

    Data is read from the database (loaded by TradeMapLoader).
    Results are exported as a three-sheet Excel workbook per partner to EXPORT_DIR.
    See README.md and docs/validation_methodology.md for the full methodology.
    """

    def __init__(self):
        self.logger = logging.getLogger(self.__class__.__name__)

        # Built during processing — keyed by partner ("China" or "US")
        self.hs6_trade_data_by_partner:   dict[str, pd.DataFrame] = {}
        self.hs4_index_by_partner:        dict[str, pd.DataFrame] = {}
        self.hs6_weighted_by_partner:     dict[str, pd.DataFrame] = {}
        self.headline_cij_by_partner:     dict[str, pd.DataFrame] = {}

    # ── Public entry point ───────────────────────────────────────────────────

    def run(self, countries: list[str] | None = None, hs4_codes: list[str] | None = None):
        self._load_from_db()
        self._filter_by_ict_scope_and_year()
        self._calculate_tci_drysdale_garnaut()
        self._calculate_rca_and_tci_rca()
        self._aggregate_hs6_to_hs4()
        self._calculate_headline_cij()
        self._export_excel(countries, hs4_codes)
        self._export_hs4_tci_charts(countries, hs4_codes)
        self._export_word_summary(countries, hs4_codes)

    # ── Load from DB ─────────────────────────────────────────────────────────

    def _load_from_db(self):
        """
        Query all four trade tables directly in long format, extract TOTAL rows
        as per-year denominators, then merge into one master DataFrame per partner.
        """
        from loadFiles.models import (
            CountryExportToPartner, CountryExportToWorld,
            PartnerImportFromWorld, WorldExport, HSProduct,
        )

        # ── Raw long-format queries ──────────────────────────────────────────

        bilateral_records = list(
            CountryExportToPartner.objects.values('reporter', 'partner', 'product_code', 'year', 'value_usd_thousands')
        )
        reporter_world_records = list(
            CountryExportToWorld.objects.values('reporter', 'product_code', 'year', 'value_usd_thousands')
        )
        partner_import_records = list(
            PartnerImportFromWorld.objects.values('partner', 'product_code', 'year', 'value_usd_thousands')
        )
        world_export_records = list(
            WorldExport.objects.values('product_code', 'year', 'value_usd_thousands')
        )

        if not world_export_records:
            self.logger.error("No WorldExport data in DB. Run TradeMapLoader first.")
            return

        hs_label_by_product_code = dict(HSProduct.objects.values_list('product_code', 'product_label'))

        bilateral_df = pd.DataFrame(bilateral_records).rename(columns={
            'reporter': 'Country',
            'product_code': 'Product code',
            'year': 'Year',
            'value_usd_thousands': 'Reporter Export To Partner',
        })

        reporter_world_df = pd.DataFrame(reporter_world_records).rename(columns={
            'reporter': 'Country',
            'product_code': 'Product code',
            'year': 'Year',
            'value_usd_thousands': 'Reporter Export To World',
        })

        partner_import_df = pd.DataFrame(partner_import_records).rename(columns={
            'product_code': 'Product code',
            'year': 'Year',
            'value_usd_thousands': 'Partner Import From World',
        })

        world_export_df = pd.DataFrame(world_export_records).rename(columns={
            'product_code': 'Product code',
            'year': 'Year',
            'value_usd_thousands': 'World Export of item k',
        })
        world_export_df['Product label'] = (
            world_export_df['Product code'].map(hs_label_by_product_code).fillna('All products')
        )

        # Ensure Year is string for consistent joining
        for df in [bilateral_df, reporter_world_df, partner_import_df, world_export_df]:
            df['Year'] = df['Year'].astype(str)

        # ── Extract TOTAL rows as denominator lookup tables ──────────────────

        def is_total(series: pd.Series) -> pd.Series:
            return series.astype(str).str.upper() == 'TOTAL'

        reporter_totals = (
            reporter_world_df[is_total(reporter_world_df['Product code'])]
            [['Country', 'Year', 'Reporter Export To World']]
            .rename(columns={'Reporter Export To World': "Reporter's Total Export To World"})
        )

        partner_totals = (
            partner_import_df[is_total(partner_import_df['Product code'])]
            [['partner', 'Year', 'Partner Import From World']]
            .rename(columns={'Partner Import From World': "Partner's Total Import From World"})
        )

        world_totals = (
            world_export_df[is_total(world_export_df['Product code'])]
            [['Year', 'World Export of item k']]
            .rename(columns={'World Export of item k': 'Total World Export'})
        )

        # ── Drop TOTAL rows from product-level data ──────────────────────────

        bilateral_df      = bilateral_df[~is_total(bilateral_df['Product code'])]
        reporter_world_df = reporter_world_df[~is_total(reporter_world_df['Product code'])]
        partner_import_df = partner_import_df[~is_total(partner_import_df['Product code'])]
        world_export_df   = world_export_df[~is_total(world_export_df['Product code'])]

        world_export_product_labels = (
            world_export_df[['Product code', 'Product label']].drop_duplicates('Product code')
        )

        # Precompute true HS4-level world totals across ALL HS6 codes in each heading.
        # This ensures the denominator in HS4 RCA/TCI is reporter-independent, even when
        # reporters have different HS6 universes (e.g. due to the HS 2007 revision of 8542).
        world_hs4_totals_df = (
            world_export_df.assign(HS4=world_export_df['Product code'].astype(str).str[:4])
            .groupby(['HS4', 'Year'], as_index=False)['World Export of item k']
            .sum()
            .rename(columns={'World Export of item k': 'World Export HS4 Total'})
        )

        # ── Build merged DataFrame per partner ──────────────────────────────

        for partner_name in ['China', 'US']:
            reporters_for_partner = bilateral_df.loc[
                bilateral_df['partner'] == partner_name, 'Country'
            ].unique()

            if len(reporters_for_partner) == 0:
                self.logger.warning("No bilateral data for partner %s in DB.", partner_name)
                continue

            # Reporter→world: keep only reporters that have bilateral trade with this partner
            reporter_world_for_partner = reporter_world_df[
                reporter_world_df['Country'].isin(reporters_for_partner)
            ].copy()

            # Bilateral exports to this partner
            bilateral_for_partner = (
                bilateral_df[bilateral_df['partner'] == partner_name]
                .drop(columns=['partner'])
            )

            # Partner imports from world (same for all reporters)
            partner_imports_for_partner = (
                partner_import_df[partner_import_df['partner'] == partner_name]
                .drop(columns=['partner'])
            )

            # Partner denominator (total imports)
            partner_total_for_partner = (
                partner_totals[partner_totals['partner'] == partner_name]
                .drop(columns=['partner'])
            )

            merged = reporter_world_for_partner.merge(
                bilateral_for_partner, on=['Country', 'Year', 'Product code'], how='left'
            )
            merged = merged.merge(
                partner_imports_for_partner, on=['Year', 'Product code'], how='left'
            )
            merged = merged.merge(
                world_export_df[['Year', 'Product code', 'World Export of item k']],
                on=['Year', 'Product code'], how='left'
            )
            merged = merged.merge(
                world_export_product_labels, on='Product code', how='left'
            )
            merged = merged.merge(
                reporter_totals, on=['Country', 'Year'], how='left'
            )
            merged = merged.merge(
                partner_total_for_partner, on=['Year'], how='left'
            )
            merged = merged.merge(
                world_totals, on=['Year'], how='left'
            )

            merged['HS4'] = merged['Product code'].astype(str).str[:4]
            merged = merged.merge(world_hs4_totals_df, on=['HS4', 'Year'], how='left')
            merged = merged.drop(columns=['HS4'])  # re-derived in _aggregate_hs6_to_hs4

            merged[['Reporter Export To Partner', 'Partner Import From World']] = (
                merged[['Reporter Export To Partner', 'Partner Import From World']].fillna(0)
            )

            merged = merged[[
                'Country', 'Year', 'Product code', 'Product label',
                'Reporter Export To Partner',
                'Reporter Export To World',
                'Partner Import From World',
                'World Export of item k',
                "Reporter's Total Export To World",
                "Partner's Total Import From World",
                'Total World Export',
                'World Export HS4 Total',
            ]]

            merged.sort_values(
                by=['Product code', 'Year'], ascending=[True, False],
                inplace=True, ignore_index=True,
            )

            self.hs6_trade_data_by_partner[partner_name] = merged
            self.logger.info("Loaded and merged %d rows for partner %s.", len(merged), partner_name)

    # ── Pipeline steps ───────────────────────────────────────────────────────

    def _filter_by_ict_scope_and_year(self):
        for partner_name, hs6_data in self.hs6_trade_data_by_partner.items():
            hs4_mask = hs6_data['Product code'].astype(str).str[:4].isin(HS4_ICT_HEADINGS)
            year_mask = hs6_data['Year'].astype(str).isin(YEARS_TO_INCLUDE)
            self.hs6_trade_data_by_partner[partner_name] = hs6_data[hs4_mask & year_mask]

    def _calculate_tci_drysdale_garnaut(self):
        """TCI = (Xi_k / Xi) × (Mj_k / Mj) × (WX / WX_k)"""
        for partner_name, hs6_data in self.hs6_trade_data_by_partner.items():
            reporter_export_share = hs6_data["Reporter Export To World"] / hs6_data["Reporter's Total Export To World"].replace(0, np.nan)
            partner_import_share  = hs6_data["Partner Import From World"] / hs6_data["Partner's Total Import From World"].replace(0, np.nan)
            inverse_world_share   = hs6_data["Total World Export"] / hs6_data["World Export of item k"].replace(0, np.nan)

            hs6_data["TCI_Drysdale_Garnaut"] = (
                reporter_export_share * partner_import_share * inverse_world_share
            ).fillna(0)

            self.hs6_trade_data_by_partner[partner_name] = hs6_data
            self.logger.info("TCI (Drysdale & Garnaut) calculated for partner %s.", partner_name)

    def _calculate_rca_and_tci_rca(self):
        """
        Balassa RCAs at HS6:
          RCA_export = (Xi_k/Xi) / (WX_k/WX)
          RCA_import = (Mj_k/Mj) / (WX_k/WX)

        Two TCI forms at HS6:
          TCI_RCA_DG_Decomposition = RCA_export × RCA_import × (WX_k/WX)
              — algebraically equal to TCI_Drysdale_Garnaut; internal consistency check.
          TCI_Tan_Fen_HS6          = RCA_export × RCA_import
              — Tan Fen (2024) form, no world-share factor; reported alongside DG.

        Active_Pair flag = 1 when reporter has world exports AND partner has world imports
        of the HS6 product. Counts the HS6 lines that contribute non-zero values to the
        TCI sum at HS4 and headline level.
        """
        for partner_name, hs6_data in self.hs6_trade_data_by_partner.items():
            world_product_share = (
                hs6_data["World Export of item k"] / hs6_data["Total World Export"].replace(0, np.nan)
            ).replace(0, np.nan)

            hs6_data["RCA Reporter Export"] = (
                (hs6_data["Reporter Export To World"] / hs6_data["Reporter's Total Export To World"].replace(0, np.nan))
                / world_product_share
            ).fillna(0)

            hs6_data["RCA Partner Import"] = (
                (hs6_data["Partner Import From World"] / hs6_data["Partner's Total Import From World"].replace(0, np.nan))
                / world_product_share
            ).fillna(0)

            hs6_data["Proportion World Trade"] = (
                hs6_data["World Export of item k"] / hs6_data["Total World Export"]
            ).replace(0, np.nan).fillna(0)

            hs6_data["TCI_RCA_DG_Decomposition"] = (
                hs6_data["RCA Reporter Export"]
                * hs6_data["RCA Partner Import"]
                * hs6_data["Proportion World Trade"]
            )

            hs6_data["TCI Tan Fen HS6"] = (
                hs6_data["RCA Reporter Export"]
                * hs6_data["RCA Partner Import"]
            )

            hs6_data["Active_Pair"] = (
                (hs6_data["Reporter Export To World"] > 0)
                & (hs6_data["Partner Import From World"] > 0)
            ).astype(int)

            self.hs6_trade_data_by_partner[partner_name] = hs6_data

        self.logger.info("HS6 RCA, DG-decomposition TCI and Tan Fen TCI calculated.")

    def _aggregate_hs6_to_hs4(self):
        """
        HS4 Cij = sum of HS6 Cij values within each HS4 heading. Preserves the
        bilateral-product matching information in HS6: a heading registers
        complementarity only where reporter and partner each have non-zero world
        flows of the same HS6 code.

        HS4 RCA = weighted average of HS6 RCA values, weight = W_k / W_HS4
        (each HS6 code's share of world trade within its HS4 heading). Algebraically
        identical to applying Balassa directly to HS4 totals; used here as auxiliary
        heading-level specialisation metric, not as input to HS4 Cij.
        """
        for partner_name, hs6_data in self.hs6_trade_data_by_partner.items():
            hs6_data = hs6_data.copy()
            hs6_data['HS4'] = hs6_data['Product code'].astype(str).str[:4]

            world_export_within_hs4 = hs6_data.groupby(
                ['Country', 'Year', 'HS4']
            )['World Export of item k'].transform('sum')
            hs6_data['World Share Within HS4'] = (
                hs6_data['World Export of item k']
                / world_export_within_hs4.replace(0, np.nan)
            ).fillna(0)

            hs6_data['RCA Export Weighted Contribution'] = (
                hs6_data['RCA Reporter Export'] * hs6_data['World Share Within HS4']
            )
            hs6_data['RCA Import Weighted Contribution'] = (
                hs6_data['RCA Partner Import'] * hs6_data['World Share Within HS4']
            )

            hs4_aggregated = (
                hs6_data.groupby(['Country', 'Year', 'HS4'])
                .agg(
                    TCI_DG_4digit         = ('TCI_Drysdale_Garnaut',             'sum'),
                    TCI_Tan_Fen_4digit    = ('TCI Tan Fen HS6',                  'sum'),
                    RCA_Export_4digit     = ('RCA Export Weighted Contribution', 'sum'),
                    RCA_Import_4digit     = ('RCA Import Weighted Contribution', 'sum'),
                    Total_Reporter_Export = ('Reporter Export To World',         'sum'),
                    Total_Partner_Import  = ('Partner Import From World',        'sum'),
                    Total_World_Export_K  = ('World Export HS4 Total',           'first'),
                    Num_Active_HS6_Pairs  = ('Active_Pair',                      'sum'),
                )
                .reset_index()
            )

            self.hs4_index_by_partner[partner_name]    = hs4_aggregated
            self.hs6_weighted_by_partner[partner_name] = hs6_data

    def _calculate_headline_cij(self):
        """
        Headline Cij for each (reporter, partner, year) — sum of HS6 Cij across the
        full ICT scope. Drysdale-Garnaut is the primary index; Tan Fen is reported
        for sensitivity comparison. This is Drysdale-Garnaut (1982, eq. 2) applied
        to the full scope of the present study.
        """
        for partner_name, hs6_data in self.hs6_trade_data_by_partner.items():
            headline = (
                hs6_data.groupby(['Country', 'Year'])
                .agg(
                    Headline_Cij_DG      = ('TCI_Drysdale_Garnaut', 'sum'),
                    Headline_Cij_Tan_Fen = ('TCI Tan Fen HS6',      'sum'),
                    Num_Active_HS6_Pairs = ('Active_Pair',          'sum'),
                )
                .reset_index()
                .sort_values(['Country', 'Year'])
                .reset_index(drop=True)
            )
            self.headline_cij_by_partner[partner_name] = headline
            self.logger.info(
                "Headline Cij computed for partner %s (%d reporter-year rows).",
                partner_name, len(headline),
            )

    # ── Export ───────────────────────────────────────────────────────────────

    def _export_excel(self, countries: list[str] | None, hs4_codes: list[str] | None):
        export_dir = Path(EXPORT_DIR)
        export_dir.mkdir(parents=True, exist_ok=True)

        country_summary_columns = {
            'Country':              'Country',
            'Year':                 'Year',
            'Headline_Cij_DG':      'Headline_Cij_Drysdale_Garnaut',
            'Headline_Cij_Tan_Fen': 'Headline_Cij_Tan_Fen',
            'Num_Active_HS6_Pairs': 'Active_HS6_Pairs',
        }

        hs4_summary_columns = {
            'Country':              'Country',
            'Year':                 'Year',
            'HS4':                  'HS4',
            'TCI_DG_4digit':        'TCI_Drysdale_Garnaut',
            'TCI_Tan_Fen_4digit':   'TCI_Tan_Fen',
            'RCA_Export_4digit':    'RCA_Reporter_Export',
            'RCA_Import_4digit':    'RCA_Partner_Import',
            'Total_Reporter_Export':'Total_Reporter_Export_USD_thousands',
            'Total_Partner_Import': 'Total_Partner_Import_USD_thousands',
            'Total_World_Export_K': 'Total_World_Export_HS4_USD_thousands',
            'Num_Active_HS6_Pairs': 'Active_HS6_Pairs',
        }

        hs6_detail_columns = {
            'Country':                          'Country',
            'Year':                             'Year',
            'HS4':                              'HS4',
            'Product code':                     'Product_Code',
            'Product label':                    'Product_Label',
            'Reporter Export To Partner':       'Reporter_Export_To_Partner_USD_thousands',
            'Reporter Export To World':         'Reporter_Export_To_World_USD_thousands',
            'Partner Import From World':        'Partner_Import_From_World_USD_thousands',
            'World Export of item k':           'World_Export_HS6_USD_thousands',
            'World Export HS4 Total':           'World_Export_HS4_Total_USD_thousands',
            "Reporter's Total Export To World": 'Total_Reporter_Export_USD_thousands',
            "Partner's Total Import From World":'Total_Partner_Import_USD_thousands',
            'Total World Export':               'Total_World_Export_USD_thousands',
            'TCI_Drysdale_Garnaut':             'TCI_Drysdale_Garnaut',
            'TCI_RCA_DG_Decomposition':         'TCI_RCA_DG_Decomposition',
            'TCI Tan Fen HS6':                  'TCI_Tan_Fen',
            'RCA Reporter Export':              'RCA_Reporter_Export',
            'RCA Partner Import':               'RCA_Partner_Import',
            'Proportion World Trade':           'Proportion_World_Trade',
            'World Share Within HS4':           'World_Share_Within_HS4',
            'Active_Pair':                      'Active_Pair',
        }

        for partner_name in self.hs4_index_by_partner:
            hs4_data = self.hs4_index_by_partner[partner_name].copy()
            hs6_data = self.hs6_weighted_by_partner[partner_name].copy()
            country_data = self.headline_cij_by_partner[partner_name].copy()

            if countries is not None:
                hs4_data = hs4_data[hs4_data['Country'].isin(countries)]
                hs6_data = hs6_data[hs6_data['Country'].isin(countries)]
                country_data = country_data[country_data['Country'].isin(countries)]

            # hs4_codes filter applies only to HS4 and HS6 sheets — Country Summary
            # always reports the full ICT-scope headline regardless of HS4 subset.
            if hs4_codes is not None:
                hs4_data = hs4_data[hs4_data['HS4'].isin(hs4_codes)]
                hs6_data = hs6_data[hs6_data['HS4'].isin(hs4_codes)]

            country_sheet = (
                country_data[list(country_summary_columns.keys())]
                .rename(columns=country_summary_columns)
                .sort_values(['Country', 'Year'])
                .reset_index(drop=True)
            )
            hs4_sheet = (
                hs4_data[list(hs4_summary_columns.keys())]
                .rename(columns=hs4_summary_columns)
                .sort_values(['Country', 'Year', 'HS4'])
                .reset_index(drop=True)
            )
            hs6_sheet = (
                hs6_data[list(hs6_detail_columns.keys())]
                .rename(columns=hs6_detail_columns)
                .sort_values(['Country', 'Year', 'HS4', 'Product_Code'])
                .reset_index(drop=True)
            )

            output_path = export_dir / f"{partner_name}_TCI.xlsx"
            with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                country_sheet.to_excel(writer, sheet_name='Country Summary', index=False)
                hs4_sheet.to_excel(writer, sheet_name='HS4 Summary', index=False)
                hs6_sheet.to_excel(writer, sheet_name='HS6 Detail', index=False)

            self.logger.info(
                "Exported %s (%d country-year rows, %d HS4 rows, %d HS6 rows).",
                output_path, len(country_sheet), len(hs4_sheet), len(hs6_sheet),
            )

    def _export_hs4_tci_charts(self, countries: list[str] | None, hs4_codes: list[str] | None):
        export_dir = Path(EXPORT_DIR)
        export_dir.mkdir(parents=True, exist_ok=True)

        for partner_name, hs4_data in self.hs4_index_by_partner.items():
            data = hs4_data.copy()

            if countries is not None:
                data = data[data['Country'].isin(countries)]
            if hs4_codes is not None:
                data = data[data['HS4'].isin(hs4_codes)]

            data = data.copy()
            data['Year'] = pd.to_numeric(data['Year'], errors='coerce')
            data = data.dropna(subset=['Year', 'TCI_DG_4digit'])
            data = data.sort_values('Year')

            for hs4_code in data['HS4'].unique():
                hs4_data_filtered = data[data['HS4'] == hs4_code]

                fig, ax = plt.subplots(figsize=(12, 6))
                for country_name in sorted(hs4_data_filtered['Country'].unique()):
                    country_rows = hs4_data_filtered[hs4_data_filtered['Country'] == country_name]
                    ax.plot(country_rows['Year'], country_rows['TCI_DG_4digit'],
                            marker='o', label=country_name)

                ax.set_title(f'Trade Complementarity Index — HS4 {hs4_code} — Partner: {partner_name}')
                ax.set_xlabel('Year')
                ax.set_ylabel('TCI (Drysdale-Garnaut, sum of HS6)')
                ax.legend(bbox_to_anchor=(1.05, 1), loc='upper left')
                ax.grid(True, linestyle='--', alpha=0.5)
                fig.tight_layout()
                fig.savefig(export_dir / f"{partner_name}_{hs4_code}_TCI_DG.png")
                plt.close(fig)

        self.logger.info("HS4 TCI charts exported.")

    def _export_word_summary(self, countries: list[str] | None, hs4_codes: list[str] | None):
        """
        Write RCA_Cij_Summary.docx — a method section plus one RCA/Cij year table
        per (reporter, HS4 heading), with the US and China partner columns merged.
        Scope follows the same countries / hs4_codes filter as the Excel export.
        Layout mirrors readings/RCA_Cij_Summary.docx.
        """
        from docx import Document

        export_dir = Path(EXPORT_DIR)
        export_dir.mkdir(parents=True, exist_ok=True)

        china = self.hs4_index_by_partner.get('China')
        us = self.hs4_index_by_partner.get('US')
        if china is None or us is None:
            self.logger.warning("Word summary skipped: need both China and US partner data.")
            return

        keep = ['Country', 'Year', 'HS4', 'TCI_DG_4digit',
                'RCA_Export_4digit', 'RCA_Import_4digit']
        china = china[keep].copy()
        us = us[keep].copy()

        if countries is not None:
            china = china[china['Country'].isin(countries)]
            us = us[us['Country'].isin(countries)]
        if hs4_codes is not None:
            china = china[china['HS4'].isin(hs4_codes)]
            us = us[us['HS4'].isin(hs4_codes)]

        merged = china.merge(
            us, on=['Country', 'Year', 'HS4'], how='outer',
            suffixes=('_CN', '_US'),
        )

        # RCA export is reporter-vs-world and partner-independent — sanity check.
        both = merged.dropna(subset=['RCA_Export_4digit_CN', 'RCA_Export_4digit_US'])
        if not both.empty:
            max_gap = (both['RCA_Export_4digit_CN'] - both['RCA_Export_4digit_US']).abs().max()
            if max_gap > 1e-6:
                self.logger.warning(
                    "RCA export differs across partners (max gap %.3g) — using China frame.",
                    max_gap,
                )

        merged['Year'] = pd.to_numeric(merged['Year'], errors='coerce')
        merged = merged.dropna(subset=['Year']).sort_values(['Country', 'HS4', 'Year'])

        document = Document()

        document.add_heading('RCA and Cij — Indo-Pacific reporters, US and China as partners, 2001–2024', level=0)

        document.add_heading('Method', level=1)

        document.add_heading('Variables', level=2)
        variables = [
            ('i',    "Reporter country (e.g. Korea, Vietnam)"),
            ('j',    "Partner country (US or China)"),
            ('k',    "An HS 6-digit product"),
            ('K',    "An HS 4-digit heading (a group of HS6 products)"),
            ('X_ik', "Country i's export of product k to the world"),
            ('X_i',  "Country i's total exports to the world"),
            ('M_jk', "Country j's import of product k from the world"),
            ('M_j',  "Country j's total imports from the world"),
            ('T_k',  "World total exports of product k"),
            ('T',    "World total exports across all products"),
        ]
        var_table = document.add_table(rows=1, cols=2)
        var_table.style = 'Light Grid Accent 1'
        var_table.rows[0].cells[0].text = 'Symbol'
        var_table.rows[0].cells[1].text = 'Meaning'
        for symbol, meaning in variables:
            row = var_table.add_row().cells
            row[0].text = symbol
            row[1].text = meaning
        document.add_paragraph('All trade values in USD thousands. Source: ITC Trade Map.')

        document.add_heading('Step 1 — Balassa RCA at HS6', level=2)
        document.add_paragraph(
            "How specialised is country i in exporting product k, compared with the world average?")
        document.add_paragraph("RCA_export = (X_ik / X_i) ÷ (T_k / T)")
        document.add_paragraph("And the same idea on the import side, for partner j:")
        document.add_paragraph("RCA_import = (M_jk / M_j) ÷ (T_k / T)")
        document.add_paragraph(
            "Above 1 means specialised in that product, below 1 means under-specialised. (Balassa 1965)")

        document.add_heading('Step 2 — Drysdale–Garnaut Cij at HS6', level=2)
        document.add_paragraph(
            "Do reporter i's exports of product k line up with partner j's imports of k?")
        document.add_paragraph("Cij = (X_ik / X_i) × (M_jk / M_j) × (T / T_k)")
        document.add_paragraph("Equivalently, using the two RCAs from Step 1:")
        document.add_paragraph("Cij = RCA_export × RCA_import × (T_k / T)")
        document.add_paragraph(
            "Both forms are computed and verified equal as an internal check. "
            "(Drysdale & Garnaut 1982)")

        document.add_heading('Step 3 — Aggregate from HS6 up to HS4', level=2)
        document.add_paragraph(
            "RCA at HS4 is computed by adding up the trade values across all HS6 codes "
            "inside the heading first, then applying the Balassa formula at the HS4 "
            "level. So the formula is the same as Step 1, just with the heading-level "
            "totals as inputs.")
        document.add_paragraph(
            "Cij at HS4 is the sum of all the HS6 Cij values inside the heading. "
            "No second weighting:")
        document.add_paragraph(
            "Cij at HS4 (heading K) = sum of Cij(k) for every HS6 code k inside K")

        document.add_heading('Step 4 — Interpretation', level=2)
        document.add_paragraph(
            "Cij above 1: reporter's export pattern matches partner's import pattern — "
            "natural trading partners (Frankel, Stein & Wei 1995).")
        document.add_paragraph("Cij near 1: world-average alignment.")
        document.add_paragraph("Cij below 1: misalignment.")
        document.add_paragraph(
            "The index is unitless. It measures POTENTIAL complementarity, not "
            "realised bilateral trade flows.")

        def fmt(value, decimals):
            if pd.isna(value):
                return ''
            return f"{value:.{decimals}f}"

        table_count = 0
        for reporter in sorted(merged['Country'].unique()):
            document.add_heading(str(reporter), level=1)
            reporter_rows = merged[merged['Country'] == reporter]

            for hs4_code in sorted(reporter_rows['HS4'].unique()):
                label = HS4_ICT_LABELS.get(hs4_code, '')
                heading = f"{reporter} — HS {hs4_code}"
                if label:
                    heading += f" ({label})"
                document.add_heading(heading, level=2)

                year_rows = reporter_rows[reporter_rows['HS4'] == hs4_code].sort_values('Year')

                table = document.add_table(rows=1, cols=6)
                table.style = 'Light Grid Accent 1'
                headers = ['Year', 'RCA export', 'RCA import (US)',
                           'RCA import (CN)', 'Cij vs US', 'Cij vs CN']
                for col, text in enumerate(headers):
                    table.rows[0].cells[col].text = text

                for _, year_row in year_rows.iterrows():
                    rca_export = year_row['RCA_Export_4digit_CN']
                    if pd.isna(rca_export):
                        rca_export = year_row['RCA_Export_4digit_US']
                    cells = table.add_row().cells
                    cells[0].text = str(int(year_row['Year']))
                    cells[1].text = fmt(rca_export, 3)
                    cells[2].text = fmt(year_row['RCA_Import_4digit_US'], 3)
                    cells[3].text = fmt(year_row['RCA_Import_4digit_CN'], 3)
                    cells[4].text = fmt(year_row['TCI_DG_4digit_US'], 4)
                    cells[5].text = fmt(year_row['TCI_DG_4digit_CN'], 4)

                table_count += 1

        output_path = export_dir / 'RCA_Cij_Summary.docx'
        document.save(output_path)
        self.logger.info(
            "Exported %s (%d reporter-HS4 tables).", output_path, table_count,
        )
